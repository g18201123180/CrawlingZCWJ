import re
from bs4 import BeautifulSoup
from docx import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import docx
from get_response import get_response


def get_law(title, url):
    response = get_response(url)
    # 用BeautifulSoup解析html
    soup = BeautifulSoup(response.read(), 'html.parser')

    ps = soup.find_all('p')

    document = Document()

    is_title = True

    for p in ps:
        style = p.get('style', None)
        if style:
            if not re.search(r'color', style):
                paragraph = document.add_paragraph()
                for pp in p:
                    run = paragraph.add_run(pp.text + "\n")

                    from docx.oxml.ns import qn

                    run.font.name = '微软雅黑'  # 注：如果想要设置中文字体，需在前面加上这一句
                    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if is_title:
                        run.bold = True  # 字体加粗
                        run.font.size = docx.shared.Pt(16)
                        is_title = False
                    else:
                        run.font.size = docx.shared.Pt(14)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 文字居中

        else:
            paragraph = document.add_paragraph()
            for pp in p:
                if pp.text and not re.search(r'京公网安备', pp.text):
                    run = paragraph.add_run(pp.text + "\n")
                    from docx.oxml.ns import qn

                    run.font.name = '宋体'  # 注：如果想要设置中文字体，需在前面加上这一句
                    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = docx.shared.Pt(10.5)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    document.save("D:\\laws\\" + title + ".docx")
